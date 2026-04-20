#!/usr/bin/env python3
"""
Build a document-grounding extraction bundle from a single PDF / DOCX / MD / TXT input.

Outputs inside bundle_dir:
- extracted.md
- extracted_meta.json
- asset_index.json
- assets/tables/*
- assets/figures/*
- assets/formulas/*

Important:
- This script does NOT generate grounded.md
- grounded.md must be written later by the agent after reading:
  - extracted.md
  - extracted_meta.json
  - asset_index.json
  - referenced assets under assets/

Version note:
- PDF: Docling-first extraction and asset export
- DOCX: Docling-first text extraction, plus python-docx/zip fallback for images and tables
- MD: minimal enhancement only — parse local markdown image references and copy them into bundle assets
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


def _slugify(text: str, max_len: int = 80) -> str:
    text = re.sub(r"\s+", "-", text.strip())
    text = re.sub(r"[^A-Za-z0-9._-]", "_", text)
    text = re.sub(r"_+", "_", text)
    text = text.strip("_-")
    return (text[:max_len] or "item")


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _write_text(path: Path, text: str) -> None:
    _ensure_dir(path.parent)
    path.write_text(text, encoding="utf-8")


def _json_safe(obj: Any) -> Any:
    if obj is None:
        return None
    if isinstance(obj, (str, int, float, bool)):
        return obj
    if isinstance(obj, Path):
        return str(obj)
    if callable(obj):
        return None
    if isinstance(obj, dict):
        return {str(k): _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple, set)):
        return [_json_safe(x) for x in obj]
    try:
        json.dumps(obj)
        return obj
    except Exception:
        try:
            return str(obj)
        except Exception:
            return None


def _safe_write_json(path: Path, data: Dict[str, Any]) -> None:
    _ensure_dir(path.parent)
    safe = _json_safe(data)
    path.write_text(json.dumps(safe, ensure_ascii=False, indent=2), encoding="utf-8")


def _clean_text(text: str) -> str:
    lines = text.splitlines()
    cleaned: List[str] = []
    blank = False
    for line in lines:
        line = line.rstrip()
        stripped = line.strip()
        if re.fullmatch(r"\d+", stripped):
            continue
        line = re.sub(r"[ \t]{2,}", " ", line)
        if stripped == "":
            if not blank:
                cleaned.append("")
            blank = True
        else:
            cleaned.append(line)
            blank = False
    while cleaned and cleaned[-1] == "":
        cleaned.pop()
    return "\n".join(cleaned).strip()


def _guess_title_from_markdown(md: str) -> Optional[str]:
    for line in md.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#"):
            return s.lstrip("#").strip()
        return s[:200]
    return None


def _extract_caption_text(obj: Any) -> Optional[str]:
    for attr in ["caption_text", "caption"]:
        value = getattr(obj, attr, None)
        if value is None:
            continue
        if isinstance(value, str):
            text = value.strip()
            if text:
                return text
        if callable(value):
            try:
                result = value()
                if isinstance(result, str) and result.strip():
                    return result.strip()
            except Exception:
                pass
        try:
            text = str(value).strip()
            if text and not text.startswith("<bound method"):
                return text
        except Exception:
            pass
    return None


def _looks_like_docling_model_root(path: Optional[str]) -> bool:
    if not path:
        return False
    p = Path(path)
    return p.exists() and p.is_dir()


def _has_rapidocr_assets(root: Optional[str]) -> bool:
    if not root:
        return False
    p = Path(root)
    required = [
        p / "RapidOcr/torch/PP-OCRv4/det/ch_PP-OCRv4_det_infer.pth",
        p / "RapidOcr/torch/PP-OCRv4/rec/ch_PP-OCRv4_rec_infer.pth",
    ]
    return all(x.exists() for x in required)


def _build_docling_converter_for_pdf() -> "DocumentConverter":
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions

    artifacts = os.getenv("DOCLING_ARTIFACTS_PATH", "").strip() or None
    pdf_opts = PdfPipelineOptions()

    if artifacts and _looks_like_docling_model_root(artifacts):
        try:
            pdf_opts.artifacts_path = artifacts
        except Exception:
            pass

    for attr, value in [
        ("generate_picture_images", True),
        ("generate_table_images", True),
        ("generate_page_images", False),
    ]:
        if hasattr(pdf_opts, attr):
            try:
                setattr(pdf_opts, attr, value)
            except Exception:
                pass

    do_ocr = _has_rapidocr_assets(artifacts)
    if hasattr(pdf_opts, "do_ocr"):
        try:
            pdf_opts.do_ocr = do_ocr
        except Exception:
            pass

    for attr in [
        "do_picture_description",
        "do_picture_classification",
        "do_formula_enrichment",
        "do_code_enrichment",
        "do_chart_extraction",
    ]:
        if hasattr(pdf_opts, attr):
            try:
                setattr(pdf_opts, attr, False)
            except Exception:
                pass

    fmt = {InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts)}
    return DocumentConverter(format_options=fmt)


def _docling_markdown_from_path(path: str) -> Tuple[str, Any]:
    from docling.document_converter import DocumentConverter
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        converter = _build_docling_converter_for_pdf()
    else:
        converter = DocumentConverter()
    result = converter.convert(path)
    doc = result.document
    md = doc.export_to_markdown()
    return md, doc


def _extract_plain_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    chunks: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])
        if rows:
            chunks.append("")
            chunks.append("DOCX table content:")
            for row in rows:
                chunks.append(" | ".join(row))
            chunks.append("")
    return "\n".join(chunks).strip()


def _extract_text_and_doc(path: str) -> Tuple[str, Any]:
    ext = Path(path).suffix.lower()
    if ext in {".pdf", ".docx", ".md", ".txt"}:
        try:
            md, doc = _docling_markdown_from_path(path)
            return md, doc
        except Exception:
            if ext == ".docx":
                return _extract_plain_docx(path), None
            if ext in {".md", ".txt"}:
                return Path(path).read_text(encoding="utf-8"), None
            raise
    raise ValueError(f"Unsupported file extension: {ext}")


def _call_table_export(table: Any, doc: Any, names: List[str]) -> Optional[Any]:
    for name in names:
        fn = getattr(table, name, None)
        if callable(fn):
            try:
                return fn(doc)
            except TypeError:
                try:
                    return fn()
                except Exception:
                    pass
            except Exception:
                pass
    return None


def _export_table_markdown(table: Any, doc: Any) -> Optional[str]:
    result = _call_table_export(table, doc, ["export_to_markdown", "to_markdown"])
    return result if isinstance(result, str) and result.strip() else None


def _export_table_html(table: Any, doc: Any) -> Optional[str]:
    result = _call_table_export(table, doc, ["export_to_html", "to_html"])
    return result if isinstance(result, str) and result.strip() else None


def _export_table_csv(table: Any, doc: Any) -> Optional[str]:
    result = _call_table_export(table, doc, ["export_to_dataframe", "to_dataframe", "export_to_df", "to_df"])
    if result is not None and hasattr(result, "to_csv"):
        try:
            return result.to_csv(index=False)
        except Exception:
            return None
    return None


def _save_pdf_table_assets(doc: Any, tables_dir: Path) -> List[Dict[str, Any]]:
    _ensure_dir(tables_dir)
    out: List[Dict[str, Any]] = []
    tables = list(getattr(doc, "tables", []) or [])
    for i, table in enumerate(tables, start=1):
        tid = f"table_{i:03d}"
        md_path = tables_dir / f"{tid}.md"
        csv_path = tables_dir / f"{tid}.csv"
        html_path = tables_dir / f"{tid}.html"
        md_text = _export_table_markdown(table, doc)
        csv_text = _export_table_csv(table, doc)
        html_text = _export_table_html(table, doc)
        if md_text:
            _write_text(md_path, md_text)
        if csv_text:
            _write_text(csv_path, csv_text)
        if html_text:
            _write_text(html_path, html_text)
        out.append({
            "id": tid,
            "md_path": str(md_path) if md_path.exists() else None,
            "csv_path": str(csv_path) if csv_path.exists() else None,
            "html_path": str(html_path) if html_path.exists() else None,
            "caption": _extract_caption_text(table),
            "source": "docling_pdf",
        })
    return out


def _call_maybe(obj: Any, method_name: str, *args) -> Any:
    fn = getattr(obj, method_name, None)
    if callable(fn):
        try:
            return fn(*args)
        except TypeError:
            try:
                return fn()
            except Exception:
                return None
        except Exception:
            return None
    return None


def _save_pil_image(image: Any, path: Path) -> bool:
    if image is None:
        return False
    if hasattr(image, "save"):
        try:
            image.save(path)
            return True
        except Exception:
            return False
    return False


def _extract_picture_image(picture: Any, doc: Any) -> Optional[Any]:
    for meth in ["get_image", "export_to_image", "to_image"]:
        img = _call_maybe(picture, meth, doc)
        if img is not None:
            return img
    for attr in ["image", "pil_image"]:
        img = getattr(picture, attr, None)
        if img is not None:
            return img
    return None


def _save_pdf_picture_assets(doc: Any, figures_dir: Path) -> List[Dict[str, Any]]:
    _ensure_dir(figures_dir)
    out: List[Dict[str, Any]] = []
    pictures = list(getattr(doc, "pictures", []) or [])
    for i, picture in enumerate(pictures, start=1):
        fid = f"figure_{i:03d}"
        img_path = figures_dir / f"{fid}.png"
        img = _extract_picture_image(picture, doc)
        ok = _save_pil_image(img, img_path)
        out.append({
            "id": fid,
            "image_path": str(img_path) if ok else None,
            "caption": _extract_caption_text(picture),
            "source": "docling_pdf",
        })
    return out


def _table_to_matrix(table: Any) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([(cell.text or "").strip() for cell in row.cells])
    return rows


def _matrix_to_markdown(matrix: List[List[str]]) -> str:
    if not matrix:
        return ""
    width = max(len(r) for r in matrix)
    padded = [r + [""] * (width - len(r)) for r in matrix]
    header = padded[0]
    body = padded[1:] if len(padded) > 1 else []
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _matrix_to_csv_text(matrix: List[List[str]]) -> str:
    import io
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in matrix:
        writer.writerow(row)
    return buf.getvalue()


def _save_docx_table_assets(docx_path: str, tables_dir: Path, start_index: int = 1) -> List[Dict[str, Any]]:
    from docx import Document
    _ensure_dir(tables_dir)
    out: List[Dict[str, Any]] = []
    doc = Document(docx_path)
    for offset, table in enumerate(doc.tables, start=0):
        tid = f"table_{start_index + offset:03d}"
        matrix = _table_to_matrix(table)
        if not matrix:
            continue
        md_path = tables_dir / f"{tid}.md"
        csv_path = tables_dir / f"{tid}.csv"
        _write_text(md_path, _matrix_to_markdown(matrix))
        _write_text(csv_path, _matrix_to_csv_text(matrix))
        out.append({
            "id": tid,
            "md_path": str(md_path),
            "csv_path": str(csv_path),
            "html_path": None,
            "caption": None,
            "source": "python_docx_fallback",
        })
    return out


def _docx_media_files(docx_path: str) -> List[Tuple[str, bytes]]:
    items: List[Tuple[str, bytes]] = []
    with zipfile.ZipFile(docx_path, "r") as zf:
        for name in zf.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                data = zf.read(name)
                items.append((name, data))
    return items


def _save_docx_picture_assets(docx_path: str, figures_dir: Path, start_index: int = 1) -> List[Dict[str, Any]]:
    _ensure_dir(figures_dir)
    out: List[Dict[str, Any]] = []
    media_items = _docx_media_files(docx_path)
    for offset, (zip_name, data) in enumerate(media_items, start=0):
        ext = Path(zip_name).suffix.lower() or ".bin"
        fid = f"figure_{start_index + offset:03d}"
        out_path = figures_dir / f"{fid}{ext}"
        out_path.write_bytes(data)
        out.append({
            "id": fid,
            "image_path": str(out_path),
            "caption": None,
            "source": "python_docx_fallback",
            "original_zip_path": zip_name,
        })
    return out


_MD_IMAGE_RE = re.compile(r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"[^"]*")?\)')


def _is_local_image_reference(ref: str) -> bool:
    if re.match(r"^[a-zA-Z]+://", ref):
        return False
    if ref.startswith("data:"):
        return False
    return True


def _save_md_referenced_images(md_path: str, raw_markdown: str, figures_dir: Path, start_index: int = 1) -> List[Dict[str, Any]]:
    _ensure_dir(figures_dir)
    out: List[Dict[str, Any]] = []
    src_dir = Path(md_path).resolve().parent
    seen: set[str] = set()

    for match in _MD_IMAGE_RE.finditer(raw_markdown):
        alt_text = (match.group(1) or "").strip()
        ref = (match.group(2) or "").strip()
        if not _is_local_image_reference(ref):
            continue
        clean_ref = ref.split("#", 1)[0].split("?", 1)[0].strip()
        if not clean_ref:
            continue
        src = (src_dir / clean_ref).resolve()
        if str(src) in seen:
            continue
        if not src.exists() or not src.is_file():
            continue
        seen.add(str(src))
        ext = src.suffix.lower() or ".bin"
        fid = f"figure_{start_index + len(out):03d}"
        dst = figures_dir / f"{fid}{ext}"
        shutil.copy2(src, dst)
        out.append({
            "id": fid,
            "image_path": str(dst),
            "caption": alt_text or None,
            "source": "markdown_reference",
            "original_path": str(src),
        })
    return out


def _extract_formula_blocks(text: str, formulas_dir: Path) -> List[Dict[str, Any]]:
    _ensure_dir(formulas_dir)
    out: List[Dict[str, Any]] = []
    patterns = [
        re.compile(r"\$\$(.+?)\$\$", re.S),
        re.compile(r"\\begin\{equation\}(.+?)\\end\{equation\}", re.S),
        re.compile(r"\\begin\{align\}(.+?)\\end\{align\}", re.S),
    ]
    seen: List[str] = []
    for pat in patterns:
        for m in pat.findall(text):
            s = m.strip()
            if len(s) < 3:
                continue
            if s in seen:
                continue
            seen.append(s)
    for i, formula in enumerate(seen, start=1):
        fid = f"formula_{i:03d}"
        path = formulas_dir / f"{fid}.txt"
        _write_text(path, formula)
        out.append({"id": fid, "path": str(path)})
    return out


def _build_asset_ref_block_for_table(entry: Dict[str, Any], bundle_dir: Path) -> str:
    rel_md = Path(entry["md_path"]).relative_to(bundle_dir) if entry.get("md_path") else None
    rel_csv = Path(entry["csv_path"]).relative_to(bundle_dir) if entry.get("csv_path") else None
    lines = ["[AssetRef]", "type: table", f"id: {entry['id']}"]
    if rel_md:
        lines.append(f"path_md: {rel_md.as_posix()}")
    if rel_csv:
        lines.append(f"path_csv: {rel_csv.as_posix()}")
    if entry.get("caption"):
        lines.append(f"summary: {entry['caption']}")
    else:
        lines.append(f"summary: exported table asset ({entry.get('source', 'unknown')})")
    lines.append("instruction: Inspect this asset before writing grounded.md if it contains key evidence, comparisons, or numeric results.")
    lines.append("[/AssetRef]")
    return "\n".join(lines)


def _build_asset_ref_block_for_figure(entry: Dict[str, Any], bundle_dir: Path) -> str:
    lines = ["[AssetRef]", "type: figure", f"id: {entry['id']}"]
    if entry.get("image_path"):
        rel = Path(entry["image_path"]).relative_to(bundle_dir)
        lines.append(f"path: {rel.as_posix()}")
    if entry.get("caption"):
        lines.append(f"summary: {entry['caption']}")
    else:
        lines.append(f"summary: exported figure asset ({entry.get('source', 'unknown')})")
    lines.append("instruction: Inspect this asset before writing grounded.md if it appears relevant to the document's main topic, claims, or comparisons.")
    lines.append("[/AssetRef]")
    return "\n".join(lines)


def _build_asset_ref_block_for_formula(entry: Dict[str, Any], bundle_dir: Path) -> str:
    rel = Path(entry["path"]).relative_to(bundle_dir)
    return "\n".join([
        "[AssetRef]",
        "type: formula",
        f"id: {entry['id']}",
        f"path: {rel.as_posix()}",
        "summary: extracted formula-like block",
        "instruction: Inspect this asset before writing grounded.md if the document's claims depend on mathematical definitions or objectives.",
        "[/AssetRef]",
    ])


def _write_extracted_markdown(bundle_dir: Path, text_md: str, asset_index: Dict[str, Any]) -> None:
    content = _clean_text(text_md)
    refs: List[str] = []
    if asset_index.get("tables"):
        for entry in asset_index["tables"]:
            refs.append(_build_asset_ref_block_for_table(entry, bundle_dir))
    if asset_index.get("figures"):
        for entry in asset_index["figures"]:
            refs.append(_build_asset_ref_block_for_figure(entry, bundle_dir))
    if asset_index.get("formulas"):
        for entry in asset_index["formulas"]:
            refs.append(_build_asset_ref_block_for_formula(entry, bundle_dir))
    if refs:
        content = content + "\n\n---\n\n## Referenced Non-Textual Assets\n\n" + "\n\n".join(refs) + "\n"
    _write_text(bundle_dir / "extracted.md", content)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a document-grounding extraction bundle.")
    parser.add_argument("input_path", help="Path to PDF / DOCX / MD / TXT")
    parser.add_argument("bundle_dir", help="Output bundle directory")
    args = parser.parse_args()

    input_path = Path(args.input_path).resolve()
    bundle_dir = Path(args.bundle_dir).resolve()
    _ensure_dir(bundle_dir)

    ext = input_path.suffix.lower()
    assets_dir = bundle_dir / "assets"
    tables_dir = assets_dir / "tables"
    figures_dir = assets_dir / "figures"
    formulas_dir = assets_dir / "formulas"
    _ensure_dir(tables_dir)
    _ensure_dir(figures_dir)
    _ensure_dir(formulas_dir)

    extracted_text, doc = _extract_text_and_doc(str(input_path))
    cleaned_text = _clean_text(extracted_text)
    guessed_title = _guess_title_from_markdown(cleaned_text)

    tables: List[Dict[str, Any]] = []
    figures: List[Dict[str, Any]] = []
    formulas: List[Dict[str, Any]] = []

    if ext == ".pdf":
        if doc is not None:
            try:
                tables = _save_pdf_table_assets(doc, tables_dir)
            except Exception:
                tables = []
            try:
                figures = _save_pdf_picture_assets(doc, figures_dir)
            except Exception:
                figures = []
    elif ext == ".docx":
        try:
            tables = _save_docx_table_assets(str(input_path), tables_dir, start_index=1)
        except Exception:
            tables = []
        try:
            figures = _save_docx_picture_assets(str(input_path), figures_dir, start_index=1)
        except Exception:
            figures = []
    elif ext == ".md":
        try:
            figures = _save_md_referenced_images(str(input_path), extracted_text, figures_dir, start_index=1)
        except Exception:
            figures = []

    try:
        formulas = _extract_formula_blocks(cleaned_text, formulas_dir)
    except Exception:
        formulas = []

    asset_index = {
        "source_file": input_path.name,
        "source_type": input_path.suffix.lstrip(".").lower(),
        "tables": tables,
        "figures": figures,
        "formulas": formulas,
    }

    _write_extracted_markdown(bundle_dir, cleaned_text, asset_index)
    _safe_write_json(bundle_dir / "asset_index.json", asset_index)

    extracted_meta = {
        "source_file": input_path.name,
        "source_type": input_path.suffix.lstrip(".").lower(),
        "original_path": str(input_path),
        "guessed_title": guessed_title,
        "char_count": len(cleaned_text),
        "line_count": len(cleaned_text.splitlines()),
        "table_count": len(tables),
        "figure_count": len(figures),
        "formula_count": len(formulas),
        "docling_artifacts_path": os.getenv("DOCLING_ARTIFACTS_PATH"),
        "grounded_note_generated_by_script": False,
        "docx_asset_fallback_used": ext == ".docx",
        "md_reference_asset_fallback_used": ext == ".md",
    }
    _safe_write_json(bundle_dir / "extracted_meta.json", extracted_meta)


if __name__ == "__main__":
    main()

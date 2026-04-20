from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\d+)[.)]\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")

# Filter patterns for report-production metadata that should not appear in exports
_HR_RE = re.compile(r"^\s*[-*_]{3,}\s*$")                # standalone --- or *** or ___
_PIPELINE_FOOTER_RE = re.compile(
    r"^\s*\*?\s*(本报告通过|报告产出：|Report produced|Manuscript prepared)",
    re.IGNORECASE,
)
_EMPTY_PARA_RE = re.compile(r"^\s*$")


def _filter_meta_lines(lines: List[str]) -> List[str]:
    """Drop horizontal rules, pipeline production footnotes, and consecutive blank lines."""
    out: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip standalone horizontal rules (---, ***, ___)
        if _HR_RE.match(line):
            # Check if the next non-blank line is a pipeline footer — if so, skip the whole block
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and _PIPELINE_FOOTER_RE.search(lines[j]):
                # Skip the HR, all following blank lines, the footer line itself, and any trailing blanks
                k = j + 1
                while k < len(lines) and not lines[k].strip():
                    k += 1
                i = k
                continue
            # Otherwise skip only the HR line itself
            i += 1
            continue

        # Skip pipeline production metadata footnotes (standalone or inline)
        if _PIPELINE_FOOTER_RE.search(line):
            i += 1
            continue

        out.append(line)
        i += 1
    return out


_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
)


# Font settings per output language
_FONT_NORMAL: Dict[str, str] = {
    "en": "Arial",
    "zh": "Noto Serif CJK SC",
}
# Fallback chain: try preferred font first, fall back to any available CJK font
_CJK_FALLBACK_FONTS: List[str] = [
    "Noto Serif CJK SC",
    "Noto Serif CJK TC",
    "Noto Sans CJK SC",
    "Noto Sans CJK TC",
    "SimSun",
    "SimHei",
    "WenQuanYi Micro Hei",
]
_FONT_MONO: str = "Courier New"


def _add_inline_markdown(paragraph, text: str, font_normal: str, font_mono: str) -> None:
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = font_mono
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            label = part[1:part.index("](")]
            run = paragraph.add_run(label)
            run.underline = True
        else:
            paragraph.add_run(part)


def _add_numbered_paragraph(doc: Document, number: str, text: str, font_normal: str, font_mono: str):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Pt(18)
    fmt.first_line_indent = Pt(-12)
    p.add_run(f"{number}. ")
    _add_inline_markdown(p, text, font_normal, font_mono)
    return p


def _is_table_block(lines: List[str]) -> bool:
    if len(lines) < 2:
        return False
    if "|" not in lines[0] or "|" not in lines[1]:
        return False
    sep = lines[1].replace("|", "").replace("-", "").replace(":", "").strip()
    return sep == ""


def _parse_table(lines: List[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for i, line in enumerate(lines):
        if i == 1:
            continue  # skip markdown separator row
        row = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(row)
    return rows


def _apply_base_styles(doc: Document, font_normal: str) -> None:
    styles = doc.styles
    styles["Normal"].font.name = font_normal
    styles["Normal"].font.size = Pt(10.5)
    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in styles:
            styles[style_name].font.name = font_normal


def _resolve_font(preferred: str) -> str:
    """Return preferred font if available, otherwise the first available fallback."""
    if shutil.which("fc-match") and shutil.which("fc-list"):
        import subprocess
        result = subprocess.run(
            ["fc-match", "-f", "%{family}\n", preferred],
            capture_output=True, text=True, timeout=5,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip().split(",")[0].strip()
    # fallback chain
    candidates = [preferred] + _CJK_FALLBACK_FONTS
    for candidate in candidates:
        if shutil.which("fc-match"):
            result = subprocess.run(
                ["fc-match", "-f", "%{family}\n", candidate],
                capture_output=True, text=True, timeout=5,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip().split(",")[0].strip()
    # last resort: just return preferred and let LibreOffice figure it out
    return preferred


def markdown_to_docx(input_report: Path, output_path: Path, output_lang: str = "en") -> Path:
    preferred = _FONT_NORMAL.get(output_lang, "Arial")
    font_normal = _resolve_font(preferred)
    font_mono = _FONT_MONO

    text = input_report.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
    lines = _filter_meta_lines(text.split("\n"))

    doc = Document()
    _apply_base_styles(doc, font_normal)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # skip consecutive blank lines
        if not line.strip():
            i += 1
            continue

        # detect table block
        if "|" in line:
            j = i
            block = []
            while j < len(lines) and lines[j].strip() and "|" in lines[j]:
                block.append(lines[j])
                j += 1
            if _is_table_block(block):
                rows = _parse_table(block)
                if rows:
                    max_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = "Table Grid"
                    for r_idx, row in enumerate(rows):
                        for c_idx in range(max_cols):
                            value = row[c_idx] if c_idx < len(row) else ""
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            _add_inline_markdown(p, value, font_normal, font_mono)
                    i = j
                    continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            p = doc.add_heading(level=level)
            _add_inline_markdown(p, heading_match.group(2).strip(), font_normal, font_mono)
            i += 1
            continue

        ordered_match = _ORDERED_RE.match(line)
        if ordered_match:
            _add_numbered_paragraph(doc, ordered_match.group(1).strip(), ordered_match.group(2).strip(), font_normal, font_mono)
            i += 1
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_markdown(p, bullet_match.group(1).strip(), font_normal, font_mono)
            i += 1
            continue

        # paragraph block
        para_lines = [line.strip()]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].rstrip()
            if not next_line.strip():
                break
            if _HEADING_RE.match(next_line) or _ORDERED_RE.match(next_line) or _BULLET_RE.match(next_line):
                break
            # table start
            if "|" in next_line:
                break
            para_lines.append(next_line.strip())
            j += 1
        p = doc.add_paragraph()
        _add_inline_markdown(p, " ".join(para_lines), font_normal, font_mono)
        i = j

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_docx(input_report: Path, output_dir: Path, output_lang: str = "en") -> Path:
    # output_dir is base_dir/format when called from export_report.py
    # -> strip /format, add /output_lang to get base_dir/lang
    final_dir = output_dir.parent / output_lang
    final_dir.mkdir(parents=True, exist_ok=True)
    output_path = final_dir / "report.docx"
    return markdown_to_docx(input_report, output_path, output_lang)
